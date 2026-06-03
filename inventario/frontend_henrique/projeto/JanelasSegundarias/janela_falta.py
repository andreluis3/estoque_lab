from .janela_base import JanelaBase

import os
import webbrowser

from docx import Document
from datetime import datetime

from PyQt6.QtCore import Qt

from PyQt6.QtWidgets import (
    QLabel,
    QPushButton,
    QMessageBox,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QScrollArea
)
from inventario.ui.theme.scrollbar import SCROLLBAR

class JanelaItensFalta(JanelaBase):
    def __init__(self, alertas):
        super().__init__("Itens em Falta")

        self.alertas = alertas

        # TEXTO SUPERIOR
        self.label = QLabel(
            f"{len(alertas)} itens estão com estoque baixo",
            self.area_conteudo
        )

        self.label.setGeometry(
            40,
            40,
            700,
            40
        )

        self.label.setStyleSheet("""
            font-size: 24px;
            font-weight: bold;
            color: white;
        """)

        # ÁREA DE SCROLL
        self.scroll = QScrollArea(self.area_conteudo)

        self.scroll.setGeometry(
            40,
            100,
            1250,
            520
        )

        self.scroll.setWidgetResizable(True)

        self.scroll.setStyleSheet(f"""
            QScrollArea {{
                background-color: transparent;
                border: none;
            }}

            QWidget {{
                background-color: transparent;
                color: white;
                border: none;
            }}

            {SCROLLBAR}
        """)

        # WIDGET INTERNO
        self.scroll_widget = QWidget()

        self.scroll_widget.setStyleSheet("""
            background-color: #0d0d0d;
            border: none;
            border-radius: 18px;
        """)

        # LAYOUT PRINCIPAL
        self.layout_lista = QVBoxLayout(self.scroll_widget)

        self.layout_lista.setContentsMargins(
            25,
            25,
            25,
            25
        )

        self.layout_lista.setSpacing(16)

        # ADICIONANDO ALERTAS
        for item in alertas:

            # CONTAINER ITEM
            item_widget = QWidget()

            item_widget.setStyleSheet("""
                background-color: #161616;
                border-radius: 14px;
            """)

            item_widget.setMinimumHeight(100)

            # LAYOUT HORIZONTAL
            layout_item = QHBoxLayout(item_widget)

            layout_item.setContentsMargins(
                20,
                15,
                20,
                15
            )

            layout_item.setSpacing(20)

            # TEXTO ITEM
            label_item = QLabel(item)

            label_item.setWordWrap(True)

            label_item.setStyleSheet("""
                font-size: 20px;
                color: #dddddd;
                background-color: transparent;
            """)

            # BOTÃO COMPRAR
            botao_comprar = QPushButton(
                "Comprar"
            )

            botao_comprar.setFixedSize(
                170,
                55
            )

            botao_comprar.setStyleSheet("""
                QPushButton {
                    background-color: #0078ff;
                    color: white;
                    font-size: 18px;
                    font-weight: bold;
                    border-radius: 14px;
                }

                QPushButton:hover {
                    background-color: #3399ff;
                }

                QPushButton:pressed {
                    background-color: #005ed1;
                }
            """)

            # FUNÇÃO BOTÃO
            botao_comprar.clicked.connect(
                lambda checked=False, texto=item:
                self.comprar_item(texto)
            )

            # ADICIONA ELEMENTOS
            layout_item.addWidget(label_item)

            layout_item.addStretch()

            layout_item.addWidget(
                botao_comprar,
                alignment=Qt.AlignmentFlag.AlignRight
            )

            # ADICIONA NA LISTA
            self.layout_lista.addWidget(item_widget)

        self.layout_lista.addStretch()

        self.scroll.setWidget(
            self.scroll_widget
        )

        # BOTÃO CANCELAR
        self.botao_cancelar = QPushButton(
            "Cancelar",
            self.container
        )

        self.botao_cancelar.setGeometry(
            920,
            760,
            180,
            65
        )

        self.botao_cancelar.setStyleSheet("""
            QPushButton {
                background-color: #1e1e1e;
                color: white;
                font-size: 20px;
                font-weight: bold;
                border-radius: 18px;
            }

            QPushButton:hover {
                background-color: #2d2d2d;
            }
        """)

        self.botao_cancelar.clicked.connect(
            self.close
        )

        # BOTÃO RELATÓRIO
        self.botao_relatorio = QPushButton(
            "Relatório",
            self
        )

        self.botao_relatorio.setGeometry(
            1130,
            760,
            180,
            65
        )

        self.botao_relatorio.setStyleSheet("""
            QPushButton {
                background-color: #0078ff;
                color: white;
                font-size: 20px;
                font-weight: bold;
                border-radius: 18px;
            }

            QPushButton:hover {
                background-color: #3399ff;
            }

            QPushButton:pressed {
                background-color: #005ed1;
            }
        """)

        self.botao_relatorio.clicked.connect(
            self.gerar_relatorio
        )

    # FUNÇÃO COMPRAR
    def comprar_item(self, nome_item):

        # REMOVE TEXTO ALERTA
        nome_limpo = nome_item.split(
            "está com apenas"
        )[0].strip()

        # FORMATA URL
        pesquisa = nome_limpo.replace(
            " ",
            "+"
        )

        # URL GOOGLE
        url = (
            f"https://www.google.com/search?q=comprar+{pesquisa}"
        )

        # ABRE NAVEGADOR
        webbrowser.open(url)

        # MENSAGEM
        QMessageBox.information(
            self,
            "Pesquisa",
            f"Pesquisando:\n\n{nome_limpo}"
        )

    # GERAR RELATÓRIO
    def gerar_relatorio(self):

        documento = Document()

        documento.add_heading(
            "Relatório de Estoque Baixo",
            level=1
        )

        data = datetime.now().strftime(
            "%d/%m/%Y %H:%M:%S"
        )

        documento.add_paragraph(
            f"Data do relatório: {data}"
        )

        documento.add_paragraph("")

        # ALERTAS
        for item in self.alertas:

            documento.add_paragraph(
                item,
                style='List Bullet'
            )

        # NOME ARQUIVO
        nome_arquivo = (
            f"relatorio_estoque_"
            f"{datetime.now().strftime('%d_%m_%Y_%H_%M')}.docx"
        )

        # DESKTOP
        desktop = os.path.join(
            os.path.expanduser("~"),
            "Desktop"
        )

        caminho_arquivo = os.path.join(
            desktop,
            nome_arquivo
        )

        # SALVA
        documento.save(caminho_arquivo)

        QMessageBox.information(
            self,
            "Relatório gerado",
            f"Relatório salvo em:\n\n{caminho_arquivo}"
        )